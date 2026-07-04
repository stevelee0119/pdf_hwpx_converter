import os
import re
import logging
import requests
import traceback
import gc

logger = logging.getLogger(__name__)

# Optional imports for Libre Engine
try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from hwpx import HwpxDocument
except ImportError:
    HwpxDocument = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    from pdf2docx import Converter
except ImportError:
    Converter = None

try:
    from deep_translator import GoogleTranslator
except ImportError:
    GoogleTranslator = None

# Optional import for win32com
try:
    import win32com.client as win32
except ImportError:
    win32 = None

def extract_filename_from_headers(headers, fallback_name):
    """Content-Disposition 헤더에서 실제 원본 파일명을 파싱합니다."""
    from urllib.parse import unquote
    cd = headers.get("Content-Disposition", "")
    if not cd:
        return fallback_name

    # 1. RFC 5987 방식 (filename*=UTF-8''utf8_filename)
    match = re.search(r"filename\*=\s*UTF-8''([^;\n]+)", cd, re.IGNORECASE)
    if match:
        try:
            return unquote(match.group(1).strip('"\''))
        except Exception:
            pass

    # 2. 일반 방식 (filename="filename" 또는 filename=filename)
    match = re.search(r'filename=\s*["\']?([^;\n"\']+)["\']?', cd, re.IGNORECASE)
    if match:
        try:
            return unquote(match.group(1).strip())
        except Exception:
            pass

    return fallback_name

# Downloader for Google Docs (DOCX) URLs
class GoogleDocsDownloader:
    """
    Google Docs URL에서 문서를 파싱하고 DOCX 형태로 로컬에 다운로드하는 클래스입니다.
    """
    @staticmethod
    def extract_document_id(url):
        # 구글 Docs URL에서 Document ID 추출하는 정규식
        # 예: https://docs.google.com/document/d/1Xxxxx_xxxx/edit
        match = re.search(r"/document/d/([a-zA-Z0-9-_]+)", url)
        if match:
            return match.group(1)
        return None

    @classmethod
    def download_as_docx(cls, url, output_path):
        doc_id = cls.extract_document_id(url)
        if not doc_id:
            raise ValueError("올바른 Google Docs URL이 아닙니다. '/document/d/문서ID' 형식이 포함되어야 합니다.")
        export_url = f"https://docs.google.com/document/d/{doc_id}/export?format=docx"
        headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"}
        # stream=True 설정으로 메모리 사용 최소화
        response = requests.get(export_url, headers=headers, stream=True, timeout=30)
        if response.status_code == 200:
            original_filename = extract_filename_from_headers(response.headers, f"gdocs_{doc_id[:12]}.docx")
            with open(output_path, 'wb') as f:
                for chunk in response.iter_content(chunk_size=32768):
                    if chunk:
                        f.write(chunk)
            return output_path, original_filename
        elif response.status_code in (401, 403):
            raise PermissionError("구글 Docs 문서에 접근할 권한이 없습니다. 링크 공유 설정을 '링크가 있는 모든 사용자에게 공개'로 변경해 주세요.")
        else:
            raise Exception(f"Google Docs 다운로드 실패 (HTTP {response.status_code}). URL을 확인해 주세요.")

# Downloader for public Google Drive PDF files
class GoogleDriveDownloader:
    """Download a PDF from a public Google Drive share link."""
    @staticmethod
    def extract_file_id(url):
        # Pattern: https://drive.google.com/file/d/<ID>/view
        match = re.search(r"/d/([a-zA-Z0-9_-]+)", url)
        if match:
            return match.group(1)
        # Alternate pattern with id parameter
        match = re.search(r"id=([a-zA-Z0-9_-]+)", url)
        if match:
            return match.group(1)
        raise ValueError("구글 드라이브 URL에서 파일 ID를 추출할 수 없습니다.")



    @staticmethod
    def _get_confirm_token(response):
        """Google Drive 대용량 파일의 바이러스 스캔 확인 토큰을 추출합니다."""
        # 쿠키에서 확인 토큰 추출 (구버전 방식)
        for key, value in response.cookies.items():
            if key.startswith('download_warning'):
                return value
        # 응답 본문에서 confirm 파라미터 추출 (신버전 방식)
        try:
            content_type = response.headers.get('Content-Type', '')
            if 'text/html' in content_type:
                body = response.content.decode('utf-8', errors='ignore')
                # confirm=t& 또는 confirm=xxxxx 패턴 검색
                match = re.search(r'confirm=([0-9A-Za-z_-]+)', body)
                if match:
                    return match.group(1)
                # uuid 패턴 검색
                match = re.search(r'"confirm":"([^"]+)"', body)
                if match:
                    return match.group(1)
        except Exception:
            pass
        return None

    @classmethod
    def download_as_pdf(cls, url, output_path):
        file_id = cls.extract_file_id(url)
        session = requests.Session()
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
        }

        # 1차 시도: 직접 다운로드
        download_url = f"https://drive.google.com/uc?export=download&id={file_id}"
        response = session.get(download_url, headers=headers, stream=True, timeout=60)

        # 대용량 파일의 경우 바이러스 스캔 확인 페이지가 반환될 수 있음
        confirm_token = cls._get_confirm_token(response)
        if confirm_token:
            print(f"[GoogleDriveDownloader] 대용량 파일 확인 토큰 감지, 재시도 중...")
            download_url = f"https://drive.google.com/uc?export=download&id={file_id}&confirm={confirm_token}"
            response = session.get(download_url, headers=headers, stream=True, timeout=120)

        # 현대적인 Google Drive 다운로드 URL 패턴도 시도
        if response.status_code != 200 or 'text/html' in response.headers.get('Content-Type', ''):
            alt_url = f"https://drive.usercontent.google.com/download?id={file_id}&export=download&authuser=0&confirm=t"
            response = session.get(alt_url, headers=headers, stream=True, timeout=120)

        if response.status_code != 200:
            raise Exception(f"Google Drive PDF 다운로드 실패 (HTTP {response.status_code}). URL을 확인해 주세요.")

        # 내용이 HTML인 경우 (접근 거부 또는 오류 페이지)
        content_type = response.headers.get('Content-Type', '')
        if 'text/html' in content_type:
            raise Exception(
                "Google Drive에서 파일을 다운로드할 수 없습니다. "
                "파일이 공개 공유 설정('링크가 있는 모든 사용자')인지 확인해 주세요."
            )

        original_filename = extract_filename_from_headers(response.headers, f"gdrive_{file_id[:12]}.pdf")

        with open(output_path, 'wb') as f:
            for chunk in response.iter_content(chunk_size=32768):
                if chunk:
                    f.write(chunk)

        # 다운로드된 파일이 실제 PDF인지 검증
        with open(output_path, 'rb') as f:
            header = f.read(4)
        if header != b'%PDF':
            os.remove(output_path)
            raise Exception(
                "다운로드된 파일이 유효한 PDF가 아닙니다. "
                "Google Drive 공유 설정을 확인하거나, 파일이 PDF 형식인지 확인해 주세요."
            )

        return output_path, original_filename

def sanitize_text_for_translation(text):
    """동일한 문자가 불필요하게 대량 반복되는 경우 번역기 폭주/버퍼 현상을 막기 위해 정규화합니다."""
    if not text:
        return ""
    # 5회 이상 연속되는 문자(알파벳, 기호 등)가 있으면 2회로 축소
    # 예: aaaaaaa -> aa, ...... -> ..
    return re.sub(r'(.)\1{4,}', r'\1\1', text)

class DocxTranslator:
    """
    DOCX 파일의 문단과 표를 순회하며 텍스트를 한국어로 번역하는 클래스입니다.
    """
    @staticmethod
    def translate_docx(docx_path, output_path, progress_callback=None):
        if GoogleTranslator is None:
            raise RuntimeError("deep-translator 라이브러리가 필요합니다. 'pip install deep-translator'를 진행해 주세요.")
        if DocxDocument is None:
            raise RuntimeError("python-docx 라이브러리가 필요합니다.")

        doc = DocxDocument(docx_path)
        translator = GoogleTranslator(source='auto', target='ko')

        # 문단 번역
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                try:
                    # 번역 전 이상 데이터(알파벳 무한반복 등) 전처리
                    sanitized_text = sanitize_text_for_translation(text)
                    translated = translator.translate(sanitized_text)
                    if translated:
                        # 기존 run들을 지우고 번역된 텍스트를 새 run으로 추가
                        for r in p.runs:
                            r.text = ''
                        p.add_run(translated)
                except Exception as e:
                    print(f"[DocxTranslator] 문단 번역 실패: {e}")

        # 표 번역
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        text = p.text.strip()
                        if text:
                            try:
                                sanitized_text = sanitize_text_for_translation(text)
                                translated = translator.translate(sanitized_text)
                                if translated:
                                    for r in p.runs:
                                        r.text = ''
                                    p.add_run(translated)
                            except Exception as e:
                                print(f"[DocxTranslator] 표 텍스트 번역 실패: {e}")
                                
        try:
            doc.save(output_path)
        finally:
            # 메모리 효율화: 가비지 컬렉션 호출
            gc.collect()


class HwpAutomationEngine:
    """
    Windows 환경에서 한컴오피스 한글 프로그램의 API(win32com)를 직접 제어하여 고품질 변환을 수행하는 클래스입니다.
    """
    @staticmethod
    def is_available():
        if win32 is None:
            return False
        try:
            # gencache 테스트, 에러 시 Dispatch fallback으로 생존력 향상
            try:
                hwp = win32.gencache.EnsureDispatch("HWPFrame.HwpObject")
            except Exception:
                hwp = win32.Dispatch("HWPFrame.HwpObject")
            hwp.Quit()
            return True
        except Exception:
            return False

    @classmethod
    def convert_file(cls, input_path, output_path):
        """
        win32com을 활용해 입력을 HWPX로 변환합니다.
        지원 입력: PDF, DOCX
        """
        if not cls.is_available():
            raise RuntimeError("한컴오피스가 설치되어 있지 않거나, 윈도우 환경이 아닙니다. OLE 자동화를 사용할 수 없습니다.")

        input_abs = os.path.abspath(input_path)
        output_abs = os.path.abspath(output_path)
        ext = os.path.splitext(input_abs)[1].lower()
        format_type = "HWPX"

        try:
            hwp = win32.gencache.EnsureDispatch("HWPFrame.HwpObject")
        except Exception:
            try:
                hwp = win32.Dispatch("HWPFrame.HwpObject")
            except Exception as e:
                raise RuntimeError(f"아래아한글 OLE 연결에 실패했습니다: {e}")

        try:
            # 보안 경고창 우회를 위한 모듈 등록 시도
            try:
                hwp.RegisterModule("FilePathCheckDLL", "SecurityModule")
            except Exception as e:
                print(f"[Automation] Security DLL register failed: {e}")

            opened = False
            if ext == ".pdf":
                # PDF 열기 시도
                try:
                    # 1순위: FileOpenPDF 액션
                    pset = hwp.HParameterSet.HFileOpenSave
                    hwp.HAction.GetDefault("FileOpenPDF", pset.HSet)
                    pset.filename = input_abs
                    result = hwp.HAction.Execute("FileOpenPDF", pset.HSet)
                    if result:
                        opened = True
                except Exception as e:
                    print(f"[Automation] FileOpenPDF Action failed: {e}, falling back to Open()")
                
                if not opened:
                    # 2순위: PDF 포맷 필터 인자를 이용한 일반 Open
                    opened = hwp.Open(input_abs, "PDF", "")
            elif ext == ".docx":
                # DOCX 열기
                opened = hwp.Open(input_abs, "docx", "")
            else:
                # 일반 문서 열기
                opened = hwp.Open(input_abs)

            if not opened:
                raise Exception(f"한글 프로그램에서 입력 파일({os.path.basename(input_path)})을 열지 못했습니다.")

            # ---- 대용량 파일 비동기 로딩 완료 대기 루프 ----
            # 아래아한글 내부적으로 대용량 파일(PDF/DOCX)을 파싱하여 렌더링을 끝낼 때까지 
            # 페이지 수(PageCount)가 서서히 늘어나므로, 안정화(페이지 수가 더 늘지 않음)될 때까지 대기합니다.
            import time
            last_page_count = 0
            stable_ticks = 0
            
            for check_idx in range(120): # 0.5초 간격으로 최대 60초 감시
                time.sleep(0.5)
                try:
                    current_page_count = hwp.PageCount
                    if current_page_count > 0:
                        if current_page_count == last_page_count:
                            stable_ticks += 1
                            # 4번 연속(2초) 페이지 수 변동 없으면 파싱이 완전히 완료된 것으로 판단
                            if stable_ticks >= 4:
                                print(f"[Automation] Document loading stabilized at {current_page_count} pages.")
                                break
                        else:
                            last_page_count = current_page_count
                            stable_ticks = 0
                except Exception:
                    pass
            else:
                print(f"[Automation] Warning: Page count stabilization timed out. Proceeding with page count: {last_page_count}")

            # HWPX 또는 HWP 형식으로 저장
            saved = False
            try:
                hwp.HAction.GetDefault("FileSaveAs_S", hwp.HParameterSet.HFileOpenSave.HSet)
                hwp.HParameterSet.HFileOpenSave.filename = output_abs
                hwp.HParameterSet.HFileOpenSave.Format = format_type # "HWPX" or "HWP"
                result = hwp.HAction.Execute("FileSaveAs_S", hwp.HParameterSet.HFileOpenSave.HSet)
                if result:
                    saved = True
            except Exception as e:
                print(f"[Automation] FileSaveAs_S action failed: {e}, falling back to SaveAs()")
 
            # 가독성 개선: 문서 전체 정렬을 '왼쪽 정렬'로, 한글/영문 줄 나눔 기준을 '어절'로 일괄 변경
            try:
                # OLE API 호출 바쁨 상태 등을 고려해 약간의 딜레이
                import time
                time.sleep(0.5)
                hwp.Run("SelectAll")
                pset = hwp.HParameterSet.HParaShape
                hwp.HAction.GetDefault("ParagraphShape", pset.HSet)
                pset.AlignType = 1        # 왼쪽 정렬 (0: 양쪽정렬, 1: 왼쪽정렬)
                pset.BreakWordHan = 1    # 한글 줄 나눔 기준 (1: 어절, 0: 글자)
                pset.BreakWordEng = 1    # 영문 줄 나눔 기준 (1: 어절, 0: 글자)
                hwp.HAction.Execute("ParagraphShape", pset.HSet)
                hwp.Run("Cancel")        # 블록 선택 해제
            except Exception as e:
                print(f"[Automation] 문단 정렬 및 줄 나눔 일괄 변경 실패: {e}")

            if not saved:
                # 예비 저장 메서드 호출
                saved = hwp.SaveAs(output_abs, format_type, "")
 
            if not saved:
                raise Exception("변환 파일 저장에 실패했습니다.")
 
        except Exception as e:
            print(f"[Automation] Error during conversion: {e}")
            traceback.print_exc()
            raise e
        finally:
            try:
                hwp.Quit()
            except Exception:
                pass


class LibBasedEngine:
    """
    한컴오피스가 없는 환경에서 오픈소스 파서를 활용해 텍스트 기반 HWPX를 직접 구축하는 클래스입니다.
    """
    @staticmethod
    def is_available():
        return HwpxDocument is not None

    @classmethod
    def _apply_global_paragraph_style(cls, doc):
        """가독성 개선: 모든 문단 양식 템플릿의 정렬을 왼쪽 정렬로, 
        영문 및 한글 줄 나눔 방식을 어절 단위(KEEP_WORD)로 일괄 변경합니다."""
        for prop_id, prop in doc.paragraph_properties.items():
            try:
                if hasattr(prop, 'align') and hasattr(prop.align, 'horizontal'):
                    prop.align.horizontal = 'LEFT'
                if hasattr(prop, 'break_setting'):
                    prop.break_setting.break_latin_word = 'KEEP_WORD'
                    prop.break_setting.break_non_latin_word = 'KEEP_WORD'
            except Exception as e:
                print(f"[LibBasedEngine] Failed to apply global para style {prop_id}: {e}")

    @classmethod
    def _extract_page_text(cls, page):
        """
        pdfplumber를 이용해 페이지 텍스트를 추출할 때 공백(띄어쓰기) 누락을 방지하는 다중 알고리즘 구현
        """
        # 1차: x_tolerance 정밀 조정을 통한 기본 텍스트 추출 시도
        orig_text = page.extract_text(x_tolerance=1.5, y_tolerance=3) or ""

        # 2차: extract_words 기반 좌표 재구성 (단어 간 띄어쓰기 강제 보장)
        try:
            words = page.extract_words(x_tolerance=1.5, y_tolerance=3)
        except Exception:
            words = []

        if not words:
            return orig_text

        # y좌표(top) 기준 그룹화 (tolerance 3pt)
        lines_dict = {}
        for word in words:
            top_key = round(word['top'] / 3.0) * 3.0
            if top_key not in lines_dict:
                lines_dict[top_key] = []
            lines_dict[top_key].append(word)

        sorted_tops = sorted(lines_dict.keys())
        reconstructed_lines = []
        for top in sorted_tops:
            line_words = sorted(lines_dict[top], key=lambda w: w['x0'])
            line_text = " ".join([w['text'] for w in line_words])
            reconstructed_lines.append(line_text)
        reconstructed_text = "\n".join(reconstructed_lines)

        # 하이픈 결합 처리: 단어가 줄바꿈으로 끊어진 경우 (예: "interna-\ntional" -> "international")
        reconstructed_text = re.sub(r'(\w+)-\s*\n\s*(\w+)', r'\1\2', reconstructed_text)
        orig_text = re.sub(r'(\w+)-\s*\n\s*(\w+)', r'\1\2', orig_text)

        # 공백 개수 비교를 통해 띄어쓰기가 더 잘 유지된 결과물 선택
        if orig_text.count(" ") >= reconstructed_text.count(" "):
            return orig_text
        return reconstructed_text

    @classmethod
    def _process_single_page(cls, pdf_path, page_index, translator=None):
        """
        단일 페이지만 열어 텍스트와 테이블을 추출한 뒤 즉시 닫습니다.
        페이지별로 PDF를 열고 닫아 메모리를 최소화합니다.
        """
        paragraphs = []
        tables_data = []

        with pdfplumber.open(pdf_path, pages=[page_index]) as pdf:
            if not pdf.pages:
                return paragraphs, tables_data
            page = pdf.pages[0]

            text = cls._extract_page_text(page)
            if text:
                paragraph_buffer = []
                for line in text.split('\n'):
                    line_str = line.strip()
                    if line_str:
                        paragraph_buffer.append(line_str)

                        if line_str.endswith((".", "?", "!", '"', "”", "’")):
                            full_para_text = " ".join(paragraph_buffer)
                            full_para_text = re.sub(r'\s+', ' ', full_para_text).strip()
                            if full_para_text:
                                if translator:
                                    try:
                                        sanitized_line = sanitize_text_for_translation(full_para_text)
                                        full_para_text = translator.translate(sanitized_line)
                                    except Exception as e:
                                        logger.warning("[LibBasedEngine] 번역 실패: %s", e)
                                paragraphs.append(full_para_text)
                            paragraph_buffer = []

                if paragraph_buffer:
                    full_para_text = " ".join(paragraph_buffer)
                    full_para_text = re.sub(r'\s+', ' ', full_para_text).strip()
                    if full_para_text:
                        if translator:
                            try:
                                sanitized_line = sanitize_text_for_translation(full_para_text)
                                full_para_text = translator.translate(sanitized_line)
                            except Exception as e:
                                logger.warning("[LibBasedEngine] 번역 실패: %s", e)
                        paragraphs.append(full_para_text)

            try:
                for table in (page.extract_tables() or []):
                    if table and len(table) > 0 and len(table[0]) > 0:
                        tables_data.append(table)
            except Exception as e:
                logger.warning("[LibBasedEngine] 테이블 추출 실패: %s", e)

        return paragraphs, tables_data

    @classmethod
    def convert_pdf_to_hwpx(cls, pdf_path, output_path, translate_to_ko=False, progress_callback=None):
        if not cls.is_available():
            raise RuntimeError("python-hwpx 라이브러리가 로드되지 않았습니다.")
        if pdfplumber is None:
            raise RuntimeError("pdfplumber 라이브러리가 필요합니다. 'pip install pdfplumber'를 진행해 주세요.")

        with pdfplumber.open(pdf_path) as pdf:
            total_pages = len(pdf.pages)

        doc = HwpxDocument.new()

        translator = None
        if translate_to_ko:
            if GoogleTranslator is None:
                logger.warning("[LibBasedEngine] deep-translator 라이브러리가 없어 번역이 지원되지 않습니다.")
            else:
                try:
                    translator = GoogleTranslator(source='auto', target='ko')
                except Exception as e:
                    logger.warning("[LibBasedEngine] 번역기 생성 실패: %s", e)

        try:
            for i in range(total_pages):
                if progress_callback:
                    pct = int(60 + (i / total_pages) * 35)
                    progress_callback(f"페이지 {i+1}/{total_pages} 처리 중...", pct)

                paragraphs, tables_data = cls._process_single_page(
                    pdf_path, i, translator
                )

                for para_text in paragraphs:
                    doc.add_paragraph(para_text)
                    doc.add_paragraph("")

                for table in tables_data:
                    rows_count = len(table)
                    cols_count = len(table[0])
                    hwp_table = doc.add_table(rows_count, cols_count)
                    for r_idx, row in enumerate(table):
                        for c_idx, cell_value in enumerate(row):
                            if cell_value is not None:
                                hwp_table.rows[r_idx].cells[c_idx].set_text(str(cell_value))

                if i < total_pages - 1:
                    doc.add_paragraph("--------------------------------------------------------------------------------")

                gc.collect()

            cls._apply_global_paragraph_style(doc)
            doc.save_to_path(output_path)
        finally:
            gc.collect()

    @classmethod
    def _docx_color_to_hex(cls, color):
        """python-docx Color 객체를 #RRGGBB 문자열로 변환합니다."""
        try:
            if color and color.type is not None and color.rgb is not None:
                return f"#{color.rgb}"
        except Exception:
            pass
        return None

    @classmethod
    def _get_run_formatting(cls, run):
        """Run의 서식 정보를 딕셔너리로 반환합니다."""
        fmt = {}
        try:
            if run.bold:
                fmt['bold'] = True
        except Exception:
            pass
        try:
            if run.italic:
                fmt['italic'] = True
        except Exception:
            pass
        try:
            if run.underline:
                fmt['underline'] = True
        except Exception:
            pass
        try:
            if run.font.name:
                fmt['font_name'] = run.font.name
        except Exception:
            pass
        try:
            if run.font.size:
                # pt 단위로 변환 (EMU → pt: 1pt = 12700 EMU)
                fmt['font_size'] = int(run.font.size / 12700)
        except Exception:
            pass
        try:
            color_hex = cls._docx_color_to_hex(run.font.color)
            if color_hex:
                fmt['color'] = color_hex
        except Exception:
            pass
        return fmt

    @classmethod
    def convert_docx_to_hwpx(cls, docx_path, output_path):
        if not cls.is_available():
            raise RuntimeError("python-hwpx 라이브러리가 로드되지 않았습니다.")
        if DocxDocument is None:
            raise RuntimeError("python-docx 라이브러리가 필요합니다. 'pip install python-docx'를 진행해 주세요.")

        doc = HwpxDocument.new()
        docx_doc = DocxDocument(docx_path)

        for para in docx_doc.paragraphs:
            # 문단 정렬 정보 추출
            alignment = None
            try:
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                align_map = {
                    WD_ALIGN_PARAGRAPH.LEFT: 'left',
                    WD_ALIGN_PARAGRAPH.CENTER: 'center',
                    WD_ALIGN_PARAGRAPH.RIGHT: 'right',
                    WD_ALIGN_PARAGRAPH.JUSTIFY: 'justify',
                }
                alignment = align_map.get(para.alignment)
            except Exception:
                pass

            # 문단 내 Run이 없거나 텍스트가 비어있는 경우 빈 줄 추가
            if not para.runs or not para.text.strip():
                doc.add_paragraph(para.text if para.text else "")
                continue

            # Run 단위로 서식 정보를 포함한 문단 구성 시도
            # python-hwpx가 run-level 서식을 지원하는 경우 사용
            run_texts = []
            has_formatting = False
            for run in para.runs:
                fmt = cls._get_run_formatting(run)
                if fmt:
                    has_formatting = True
                run_texts.append((run.text, fmt))

            if has_formatting:
                # 서식이 있는 경우: add_paragraph에 전체 텍스트 추가 후 서식 시도
                # python-hwpx의 API 수준에 따라 적용 가능 여부가 다름
                full_text = para.text

                # 가장 많이 사용된 Run의 서식을 대표 서식으로 사용
                representative_fmt = {}
                if run_texts:
                    # 텍스트 길이 기준으로 가장 긴 Run의 서식을 사용
                    longest_run = max(run_texts, key=lambda x: len(x[0]))
                    representative_fmt = longest_run[1]

                try:
                    p = doc.add_paragraph(full_text)
                    # 서식 적용 시도 (python-hwpx API가 지원하는 경우)
                    if hasattr(p, 'set_bold') and representative_fmt.get('bold'):
                        p.set_bold(True)
                    if hasattr(p, 'set_italic') and representative_fmt.get('italic'):
                        p.set_italic(True)
                    if hasattr(p, 'set_underline') and representative_fmt.get('underline'):
                        p.set_underline(True)
                    if hasattr(p, 'set_font_size') and representative_fmt.get('font_size'):
                        p.set_font_size(representative_fmt['font_size'])
                    if hasattr(p, 'set_font_name') and representative_fmt.get('font_name'):
                        p.set_font_name(representative_fmt['font_name'])
                    if hasattr(p, 'set_text_color') and representative_fmt.get('color'):
                        p.set_text_color(representative_fmt['color'])
                    if hasattr(p, 'set_alignment') and alignment:
                        p.set_alignment(alignment)
                except Exception as e:
                    # add_paragraph가 단순히 문자열을 기대하는 경우 fallback
                    print(f"[LibBasedEngine] 서식 적용 실패 (무시): {e}")
            else:
                doc.add_paragraph(para.text)

            # 가독성 개선: 실제 내용이 채워진 문단 뒤에만 띄어쓰기용 빈 문단 추가
            if para.text.strip():
                doc.add_paragraph("")

        # 표(Table) 처리
        for table in docx_doc.tables:
            try:
                rows_count = len(table.rows)
                if rows_count == 0:
                    continue
                cols_count = len(table.rows[0].cells)
                if cols_count == 0:
                    continue
                
                hwp_table = doc.add_table(rows_count, cols_count)
                for r_idx, row in enumerate(table.rows):
                    for c_idx, cell in enumerate(row.cells):
                        if r_idx < rows_count and c_idx < cols_count:
                            hwp_table.rows[r_idx].cells[c_idx].set_text(cell.text)
            except Exception as e:
                print(f"[LibBasedEngine] 테이블 변환 실패 (무시): {e}")

        # 가독성 개선: 모든 테이블 변환 완료 후 전역 문단 스타일 설정 적용
        cls._apply_global_paragraph_style(doc)

        doc.save_to_path(output_path)


class UniversalConverter:
    """
    엔진 선택 및 입력을 조합하여 전체 변환 흐름을 중재하는 메인 컨트롤러 클래스입니다.
    """
    @classmethod
    def convert(cls, input_source, output_path, use_automation=True, translate_to_ko=False, progress_callback=None):
        """
        input_source: 구글 Docs URL, Google Drive URL 또는 로컬 파일 경로 (PDF, DOCX)
        output_path: 변환 결과가 저장될 HWPX 파일 경로 (기본 저장 디렉토리 + 기본 파일명)
        use_automation: True 시 한글 OLE 자동화 사용, False 시 오픈소스 라이브러리 사용
        translate_to_ko: True 시 영어를 한국어로 자동 번역
        """
        temp_files = []
        try:
            # 1단계: 구글 Docs URL 확인 및 다운로드
            is_url = input_source.startswith("http://") or input_source.startswith("https://")
            local_input = input_source
            original_filename = None  # 원본 파일명 (파일명 생성에 사용)

            if is_url:
                if progress_callback:
                    progress_callback("URL에서 파일 다운로드 중...", 20)

                if "drive.google.com" in input_source:
                    # Google Drive: 임시 파일로 PDF 다운로드
                    temp_pdf = output_path + ".temp.pdf"
                    temp_pdf, original_filename = GoogleDriveDownloader.download_as_pdf(input_source, temp_pdf)
                    temp_files.append(temp_pdf)
                    local_input = temp_pdf
                else:
                    # Google Docs: DOCX로 다운로드
                    temp_docx = output_path + ".temp.docx"
                    temp_docx, original_filename = GoogleDocsDownloader.download_as_docx(input_source, temp_docx)
                    temp_files.append(temp_docx)
                    local_input = temp_docx
            else:
                # 로컬 파일: 원본 파일명에서 기본 이름 추출
                original_filename = os.path.splitext(os.path.basename(input_source))[0]

            # 출력 경로를 원본 파일명 기반으로 재구성
            if original_filename:
                # 다운로드한 파일명에서 base 추출
                base_name = os.path.splitext(original_filename)[0]
                save_dir = os.path.dirname(output_path)
                safe_base = re.sub(r'[<>:"/\\|?*]', '_', base_name)
                output_path = os.path.join(save_dir, f"{safe_base}(converted).hwpx")
                print(f"[UniversalConverter] 원본 파일명 기반 출력 경로 결정: {output_path}")

            if progress_callback:
                progress_callback("파일 변환 준비 중...", 40)

            # 2단계: 변환 방식 선택 및 실행
            ext = os.path.splitext(local_input)[1].lower()

            if use_automation:
                if not HwpAutomationEngine.is_available():
                    raise RuntimeError("한컴오피스 자동화 엔진을 사용할 수 없는 환경입니다. (Windows OS & 한컴오피스 필수)")

                # OLE 자동화를 사용하는 경우:
                if ext == ".pdf":
                    # PDF 파일이면 pdf2docx 사전 변환을 거치지 않고, 
                    # 아래아한글 프로그램의 강력한 내장 PDF 변환기를 직접 사용하도록 우선 시도합니다.
                    if progress_callback:
                        progress_callback("한컴오피스 PDF 변환 필터로 문서 열기 시도 중...", 60)
                    try:
                        HwpAutomationEngine.convert_file(local_input, output_path)
                        # 성공 시 바로 완료 리턴
                        if progress_callback:
                            progress_callback(f"변환 완료! 결과 파일: {os.path.basename(output_path)}", 100)
                        return output_path
                    except Exception as e:
                        print(f"[UniversalConverter] OLE 직접 PDF 로드 실패: {e}. pdf2docx 중간 변환 방식으로 우회합니다.")
                        # 한글의 PDF 직접 로딩이 오류날 때만 pdf2docx -> docx -> 한글 OLE 로드 순으로 fallback 진행
                        if Converter is None:
                            raise RuntimeError("pdf2docx 라이브러리가 필요합니다. 'pip install pdf2docx'를 진행해 주세요.")
                        if progress_callback:
                            progress_callback("PDF 레이아웃 호환 변환을 위한 DOCX 사전 변환 중...", 65)
                        temp_pdf_docx = output_path + ".converted.docx"
                        cv = Converter(local_input)
                        cv.convert(temp_pdf_docx, start=0, end=None)
                        cv.close()
                        temp_files.append(temp_pdf_docx)
                        local_input = temp_pdf_docx
                        ext = ".docx"

                # 번역 옵션 적용 (DOCX 변환 상태에서만 지원되므로 순수 DOCX에 대해 번역 후 자동화 호출)
                if translate_to_ko and ext == ".docx":
                    if progress_callback:
                        progress_callback("문서 내용을 한국어로 번역 중입니다...", 75)
                    temp_translated = output_path + ".translated.docx"
                    DocxTranslator.translate_docx(local_input, temp_translated, progress_callback)
                    temp_files.append(temp_translated)
                    local_input = temp_translated

                if progress_callback:
                    progress_callback("한컴오피스 연동 엔진으로 문서 변환 중...", 85)
                HwpAutomationEngine.convert_file(local_input, output_path)

            else:
                # 오픈소스 엔진을 사용하는 경우:
                if progress_callback:
                    progress_callback("오픈소스 엔진으로 레이아웃 파싱 및 HWPX 생성 중...", 60)

                if ext == ".pdf":
                    LibBasedEngine.convert_pdf_to_hwpx(local_input, output_path, translate_to_ko, progress_callback)
                elif ext == ".docx":
                    # DOCX 파일이면 텍스트 및 스타일 파싱
                    if translate_to_ko:
                        if progress_callback:
                            progress_callback("문서 번역 중...", 70)
                        temp_translated = output_path + ".translated.docx"
                        DocxTranslator.translate_docx(local_input, temp_translated, progress_callback)
                        temp_files.append(temp_translated)
                        local_input = temp_translated
                    LibBasedEngine.convert_docx_to_hwpx(local_input, output_path)
                else:
                    raise ValueError(f"지원하지 않는 입력 형식입니다: {ext}")

            if progress_callback:
                progress_callback(f"변환 완료! 결과 파일: {os.path.basename(output_path)}", 100)

            return output_path

        finally:
            # 임시 파일 정리
            for temp_file in temp_files:
                if os.path.exists(temp_file):
                    try:
                        os.remove(temp_file)
                    except Exception as e:
                        print(f"Warning: Failed to delete temp file {temp_file}: {e}")
            # 메모리 효율화: 가비지 컬렉션 호출
            gc.collect()
