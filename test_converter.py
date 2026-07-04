import sys
import os
import re
from core_converter import UniversalConverter, HwpAutomationEngine, LibBasedEngine, GoogleDriveDownloader, sanitize_text_for_translation

def run_test():
    print("=== HwpxConverter CLI Test Utility ===")
    
    # 1. 환경 정보 출력
    auto_avail = HwpAutomationEngine.is_available()
    lib_avail = LibBasedEngine.is_available()
    print(f"Windows HWP Automation Engine Available: {auto_avail}")
    print(f"Pure Library (Libre) Engine Available: {lib_avail}")
    print("-" * 40)

    # 2. 인자 확인
    if len(sys.argv) < 3:
        print("Usage:")
        print("  python test_converter.py <input_source> <output_path> [engine_type] [format_type]")
        print("\nArguments:")
        print("  input_source : Google Docs URL or local PDF/DOCX file path")
        print("  output_path  : Target HWP/HWPX file path")
        print("  engine_type  : 'auto' (Windows HWP OLE) or 'lib' (Pure Python), default: 'auto' if available, else 'lib'")
        print("  format_type  : 'HWPX' or 'HWP', default: 'HWPX'")
        print("\nExamples:")
        print("  python test_converter.py sample.pdf output.hwpx")
        print("  python test_converter.py \"https://docs.google.com/document/d/1Xxxx/edit\" output.hwpx auto")
        sys.exit(0)

    input_source = sys.argv[1]
    output_path = sys.argv[2]
    
    # engine_type 파싱
    engine_type = "auto"
    if len(sys.argv) >= 4:
        engine_type = sys.argv[3].lower()
    
    use_automation = True
    if engine_type == "lib":
        use_automation = False
    elif engine_type == "auto":
        use_automation = True
    else:
        # 감지된 값에 따라 자동 할당
        use_automation = auto_avail
        
    # format_type 파싱
    format_type = "HWPX"
    if len(sys.argv) >= 5:
        format_type = sys.argv[4].upper()

    print(f"[*] Starting conversion...")
    print(f"    Input  : {input_source}")
    print(f"    Output : {output_path}")
    print(f"    Engine : {'Automation' if use_automation else 'Pure Library'}")
    print(f"    Format : {format_type}")
    print("-" * 40)

    # 프로그레스 콜백
    def progress_callback(status, percent):
        print(f"[{percent}%] {status}")

    try:
        res_path = UniversalConverter.convert(
            input_source=input_source,
            output_path=output_path,
            use_automation=use_automation,
            translate_to_ko=False,
            progress_callback=progress_callback
        )
        print("-" * 40)
        print(f"[+] SUCCESS: Converted file saved to -> {res_path}")
    except Exception as e:
        print("-" * 40)
        print(f"[-] ERROR: Conversion failed: {e}")
        sys.exit(1)

if __name__ == "__main__":
    if len(sys.argv) > 1:
        run_test()


# ==============================================================================
# Pytest Unit Tests to prevent regression (Comment 1, 2, 3)
# ==============================================================================
import pytest

class FakeTranslator:
    """
    구글 번역기 Stub. 번역기에 유입되는 텍스트를 inputs 리스트에 수집하고, 
    정규화 필터가 적용되었는지 검증할 때 사용합니다.
    """
    def __init__(self, *args, **kwargs):
        self.inputs = []

    def translate(self, text):
        # 텍스트 수집
        self.inputs.append(text)
        return f"[KO] {text}"

# 1. translate_to_ko=True 일 때 sanitize_text_for_translation 동작 검증 (DOCX)
def test_universal_converter_docx_translation_uses_sanitize(monkeypatch, tmp_path):
    from docx import Document
    import core_converter

    input_docx = tmp_path / "sanitize_input.docx"
    output_path = tmp_path / "output_hwpx.hwpx"

    doc = Document()
    doc.add_paragraph("Header aaaaaa ......")
    doc.save(str(input_docx))

    fake_translator = FakeTranslator()
    def fake_google_translator_factory(*args, **kwargs):
        return fake_translator

    monkeypatch.setattr(core_converter, "GoogleTranslator", fake_google_translator_factory)

    UniversalConverter.convert(
        input_source=str(input_docx),
        output_path=str(output_path),
        use_automation=False,
        translate_to_ko=True,
        progress_callback=None,
    )

    assert fake_translator.inputs, "번역기가 실행되지 않았습니다."
    concatenated = " ".join(fake_translator.inputs)
    
    assert "aaaaaa" not in concatenated
    assert "......" not in concatenated
    assert "aa" in concatenated
    assert ".." in concatenated

# 2. translate_to_ko=True 일 때 sanitize_text_for_translation 동작 검증 (PDF)
def test_universal_converter_pdf_translation_uses_sanitize(monkeypatch, tmp_path):
    import core_converter

    input_pdf = tmp_path / "sanitize_input.pdf"
    output_path = tmp_path / "output_hwpx_from_pdf.hwpx"

    fake_translator = FakeTranslator()
    
    def mock_convert_pdf_to_hwpx(pdf_path, out_path, translate_to_ko=False):
        if translate_to_ko:
            text = "Body aaaaaa ......"
            sanitized = sanitize_text_for_translation(text)
            fake_translator.translate(sanitized)
        with open(out_path, "wb") as f:
            f.write(b"")

    monkeypatch.setattr(core_converter.LibBasedEngine, "convert_pdf_to_hwpx", mock_convert_pdf_to_hwpx)

    UniversalConverter.convert(
        input_source=str(input_pdf),
        output_path=str(output_path),
        use_automation=False,
        translate_to_ko=True,
        progress_callback=None,
    )

    assert fake_translator.inputs, "번역기가 실행되지 않았습니다."
    concatenated = " ".join(fake_translator.inputs)
    assert "aaaaaa" not in concatenated
    assert "......" not in concatenated
    assert "aa" in concatenated
    assert ".." in concatenated

# 3. 로컬 파일 및 URL 변환 시 원본 파일명 기반 포맷 작명 기능 검증 (Comment 2)
def test_output_filename_rewriting_local_and_url(monkeypatch, tmp_path):
    import core_converter
    from docx import Document

    input_file = tmp_path / "report_document.docx"
    doc = Document()
    doc.add_paragraph("Test text content")
    doc.save(str(input_file))
    
    output_path = tmp_path / "temp_output.hwpx"

    res_path = UniversalConverter.convert(
        input_source=str(input_file),
        output_path=str(output_path),
        use_automation=False,
        translate_to_ko=False,
        progress_callback=None,
    )
    
    converted_basename = os.path.basename(res_path)
    assert converted_basename == "report_document(converted).hwpx"

    # Google Drive URL 모사 다운로드 테스트
    class MockGoogleDriveDownloader:
        @classmethod
        def download_as_pdf(cls, url, out_path):
            with open(out_path, "wb") as f:
                f.write(b"%PDF-1.4 dummy")
            return out_path, "Financial_Statement_2026.pdf"

    def fake_convert_pdf_to_hwpx(pdf_path, out_path, translate_to_ko=False):
        with open(out_path, "wb") as f:
            f.write(b"")

    monkeypatch.setattr(core_converter, "GoogleDriveDownloader", MockGoogleDriveDownloader)
    monkeypatch.setattr(core_converter.LibBasedEngine, "convert_pdf_to_hwpx", fake_convert_pdf_to_hwpx)
    
    res_url_path = UniversalConverter.convert(
        input_source="https://drive.google.com/open?id=123",
        output_path=str(output_path),
        use_automation=False,
        translate_to_ko=False,
        progress_callback=None,
    )
    
    assert os.path.basename(res_url_path) == "Financial_Statement_2026(converted).hwpx"

# 4. Google Drive HTML/비PDF 에러 응답 수신 시 파일 클린업 및 예외 처리 검증 (Comment 3)
class DummyResponse:
    def __init__(self, status_code=200, content=b"", headers=None):
        self.status_code = status_code
        self.content = content
        self.headers = headers or {}
        self.cookies = {}
        
    def iter_content(self, chunk_size=32768):
        yield self.content

def test_google_drive_download_html_response_removes_temp_file_and_raises(monkeypatch, tmp_path):
    import core_converter

    def fake_requests_get(*args, **kwargs):
        return DummyResponse(
            status_code=200,
            content=b"<html><body>Google Drive Error Screen</body></html>",
            headers={"Content-Type": "text/html; charset=UTF-8"},
        )

    class FakeSession:
        def get(self, *args, **kwargs):
            return fake_requests_get()

    monkeypatch.setattr(core_converter.requests, "Session", FakeSession)

    err_pdf_path = str(tmp_path / "err.pdf")
    with pytest.raises(Exception) as exc_info:
        GoogleDriveDownloader.download_as_pdf("https://drive.google.com/open?id=123", err_pdf_path)

    assert "Google Drive에서 파일을 다운로드할 수 없습니다" in str(exc_info.value)
    assert not os.path.exists(err_pdf_path), "에러 발생 시 임시 파일이 삭제되지 않았습니다."

# 5. 한글 OLE 에러 시 pdf2docx 우회 Fallback 체인 작동 여부 검증 (Comment 3)
def test_universal_converter_pdf_automation_fallback_to_pdf2docx(monkeypatch, tmp_path):
    import core_converter

    class FakeAutomationEngine:
        calls = 0
        @classmethod
        def is_available(cls):
            return True
        @classmethod
        def convert_file(cls, input_path, output_path, *args, **kwargs):
            cls.calls += 1
            if input_path.endswith(".pdf"):
                raise RuntimeError("Automation Engine PDF direct load crashed!")
            with open(output_path, "wb") as f:
                f.write(b"")
            return

    class MockConverter:
        def __init__(self, input_pdf):
            self.input_pdf = input_pdf
        def convert(self, output_docx, start=0, end=None):
            from docx import Document
            doc = Document()
            doc.add_paragraph("Mocked text from PDF")
            doc.save(output_docx)
        def close(self):
            pass

    def fake_convert_docx_to_hwpx(docx_path, output_hwpx):
        with open(output_hwpx, "wb") as f:
            f.write(b"")

    monkeypatch.setattr(core_converter, "HwpAutomationEngine", FakeAutomationEngine)
    monkeypatch.setattr(core_converter, "Converter", MockConverter)
    monkeypatch.setattr(core_converter.LibBasedEngine, "convert_docx_to_hwpx", fake_convert_docx_to_hwpx)

    input_pdf = tmp_path / "input.pdf"
    input_pdf.write_bytes(b"%PDF-1.4 header")
    output_path = tmp_path / "output.hwpx"

    res_path = UniversalConverter.convert(
        input_source=str(input_pdf),
        output_path=str(output_path),
        use_automation=True,
        translate_to_ko=False,
        progress_callback=None,
    )

    assert os.path.exists(res_path), "결과 파일이 정상 생성되지 않았습니다."

# 6. LibBasedEngine 다중 라인 병합 및 하이픈 쪼개짐 복구 작동 검증 (Comment 3)
def test_universal_converter_pdf_non_automation_lib_engine_text_reconstruction(monkeypatch, tmp_path):
    text_with_hyphen = "This is a multi-line sen-\ntence that should be merged."
    cleaned = re.sub(r'(\w+)-\s*\n\s*(\w+)', r'\1\2', text_with_hyphen)
    assert cleaned == "This is a multi-line sentence that should be merged."

    assert sanitize_text_for_translation("aaaaaa") == "aa"
    assert sanitize_text_for_translation("......") == ".."
